"""Export Course content to a Word (.docx) document."""

from __future__ import annotations

import logging
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from ..models import ContentBlock, ContentBlockType, Course

log = logging.getLogger(__name__)

MAX_IMAGE_WIDTH = Inches(5.5)


class DocxExporter:
    def __init__(self, output_path: Path):
        self.output_path = output_path
        self.doc = Document()
        self._setup_styles()

    def _setup_styles(self):
        """Configure base document styles."""
        style = self.doc.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(11)

        # Adjust heading colors
        for level in range(1, 5):
            name = f"Heading {level}"
            if name in self.doc.styles:
                self.doc.styles[name].font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

    def export(self, course: Course) -> Path:
        """Convert a Course into a Word document."""
        # Title page
        title_para = self.doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_para.add_run(course.title)
        run.font.size = Pt(28)
        run.bold = True
        run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

        self.doc.add_paragraph()  # spacer

        meta = self.doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta.add_run(f"Source: {course.source_url}\n").font.size = Pt(9)
        meta.add_run(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}").font.size = Pt(9)

        # Table of contents
        self.doc.add_page_break()
        self.doc.add_heading("Table of Contents", level=1)
        for i, module in enumerate(course.modules, 1):
            toc_para = self.doc.add_paragraph()
            run = toc_para.add_run(f"{i}. {module.title}")
            run.bold = True
            if module.duration_minutes:
                run = toc_para.add_run(f"  ({module.duration_minutes} min)")
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

            for j, unit in enumerate(module.units, 1):
                unit_para = self.doc.add_paragraph()
                unit_para.paragraph_format.left_indent = Inches(0.5)
                unit_para.add_run(f"{i}.{j} {unit.title}")
                if unit.duration_minutes:
                    r = unit_para.add_run(f"  ({unit.duration_minutes} min)")
                    r.font.size = Pt(9)
                    r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

        # Content
        for module in course.modules:
            self.doc.add_page_break()
            self.doc.add_heading(module.title, level=1)

            if module.summary:
                p = self.doc.add_paragraph()
                run = p.add_run(module.summary)
                run.italic = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

            for unit in module.units:
                self.doc.add_heading(unit.title, level=2)

                if unit.duration_minutes:
                    dur = self.doc.add_paragraph()
                    r = dur.add_run(f"⏱ {unit.duration_minutes} minutes")
                    r.font.size = Pt(9)
                    r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

                for block in unit.content_blocks:
                    self._add_content_block(block)

        self.doc.save(str(self.output_path))
        log.info("Word document saved: %s", self.output_path)
        return self.output_path

    def _add_content_block(self, block: ContentBlock):
        match block.block_type:
            case ContentBlockType.HEADING:
                # Map original level: h2→3, h3→4 (since module=1, unit=2)
                doc_level = min(block.level + 1, 4)
                self.doc.add_heading(block.text, level=doc_level)

            case ContentBlockType.PARAGRAPH:
                self.doc.add_paragraph(block.text)

            case ContentBlockType.LIST:
                for item in block.list_items:
                    style = "List Number" if block.ordered else "List Bullet"
                    self.doc.add_paragraph(item, style=style)

            case ContentBlockType.TABLE:
                self._add_table(block)

            case ContentBlockType.IMAGE:
                self._add_image(block)

            case ContentBlockType.NOTE:
                self._add_note(block)

            case ContentBlockType.CODE:
                p = self.doc.add_paragraph()
                run = p.add_run(block.text)
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                p.paragraph_format.left_indent = Inches(0.3)

            case ContentBlockType.QUIZ:
                self._add_quiz(block)

    def _add_image(self, block: ContentBlock):
        """Add an image to the document."""
        if block.image_path and block.image_path.exists():
            try:
                self.doc.add_picture(str(block.image_path), width=MAX_IMAGE_WIDTH)
                # Center the image
                last_para = self.doc.paragraphs[-1]
                last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception as e:
                log.warning("Failed to add image %s: %s", block.image_path, e)
                self.doc.add_paragraph(f"[Image: {block.image_alt}]")
        else:
            if block.image_alt:
                p = self.doc.add_paragraph()
                run = p.add_run(f"[Image: {block.image_alt}]")
                run.italic = True
                run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        # Caption
        if block.image_alt:
            caption = self.doc.add_paragraph()
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = caption.add_run(block.image_alt)
            run.italic = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    def _add_table(self, block: ContentBlock):
        """Add a table with styled header row."""
        cols = len(block.table_headers) if block.table_headers else (
            len(block.table_rows[0]) if block.table_rows else 0
        )
        if cols == 0:
            return

        num_rows = (1 if block.table_headers else 0) + len(block.table_rows)
        table = self.doc.add_table(rows=num_rows, cols=cols)
        table.style = "Light Grid Accent 1"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        row_idx = 0
        if block.table_headers:
            for j, header in enumerate(block.table_headers):
                cell = table.cell(0, j)
                cell.text = header
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
            row_idx = 1

        for row_data in block.table_rows:
            for j, cell_text in enumerate(row_data):
                if j < cols:
                    table.cell(row_idx, j).text = cell_text
            row_idx += 1

        self.doc.add_paragraph()  # spacer after table

    def _add_note(self, block: ContentBlock):
        """Add a callout note with visual distinction."""
        prefixes = {
            "tip": "💡 Tip",
            "note": "📝 Note",
            "warning": "⚠️ Warning",
            "important": "❗ Important",
        }
        prefix = prefixes.get(block.note_type, "Note")

        p = self.doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        run = p.add_run(f"{prefix}: ")
        run.bold = True
        run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)
        p.add_run(block.text)

    def _add_quiz(self, block: ContentBlock):
        """Add a knowledge check question."""
        p = self.doc.add_paragraph()
        run = p.add_run(block.quiz_question)
        run.bold = True

        for i, option in enumerate(block.quiz_options):
            letter = chr(65 + i)  # A, B, C, ...
            opt_p = self.doc.add_paragraph()
            opt_p.paragraph_format.left_indent = Inches(0.5)
            opt_p.add_run(f"{letter}. {option}")
